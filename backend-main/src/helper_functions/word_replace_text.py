from docx import Document


def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)


def main_replace_text(pattern_path, result_path, texts: dict):
    doc = Document(pattern_path)
    for key, value in texts.items():
        replace_text(doc, key, value)

    doc.save(result_path)
    return result_path
